# ========== LOAD MODULES ========================
import streamlit as st
import os
import json
import pandas as pd
import PyPDF2
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Langchain & AI Modules
from langchain_google_genai import ChatGoogleGenerativeAI
from tavily import TavilyClient

# ========== PAGE CONFIGURATION ==================
st.set_page_config(
    page_title="Career Engine",
    page_icon="🧭",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========== SESSION STATE INIT ==================
if "base_resume" not in st.session_state:
    st.session_state.base_resume = None
if "jobs_data" not in st.session_state:
    st.session_state.jobs_data = []
if "tailored_resumes" not in st.session_state:
    st.session_state.tailored_resumes = {}

# ========== SIDEBAR & THEME SETUP ===============
with st.sidebar:
    st.markdown(
        "<div style='display:flex; align-items:center; gap:8px; margin-bottom:2px;'>"
        "<div style='width:9px; height:9px; background:#2D4A3E; border-radius:2px;'></div>"
        "<span style='font-family:\"Fraunces\",serif; font-weight:600; font-size:1.08rem;'>Career Engine</span></div>"
        "<div style='font-size:0.72rem; color:gray; margin:0 0 22px 17px;'>Your application, in one line</div>",
        unsafe_allow_html=True
    )
    
    st.markdown("### 🎨 Theme")
    theme_choice = st.radio("Appearance", ["Light Mode ☀️", "Dark Mode 🌙"], horizontal=True, label_visibility="collapsed")
    is_dark = "Dark" in theme_choice

# ========== DESIGN SYSTEM (CSS) =================
if is_dark:
    theme_vars = """
    --paper:#0f172a; --paper-raised:#1e293b; --ink:#f8fafc; --ink-soft:#94a3b8;
    --forest:#38bdf8; --forest-dim:#0284c7; --gold:#f59e0b; --gold-soft:#78350f;
    --sage:rgba(56,189,248,0.1); --sage-line:rgba(56,189,248,0.2); --clay:#ef4444; --clay-soft:rgba(239,68,68,0.1);
    --rail:#334155; --shadow:0 4px 16px rgba(0,0,0,0.3);
    """
else:
    theme_vars = """
    --paper:#F7F5F0; --paper-raised:#FCFBF8; --ink:#1A1918; --ink-soft:#5B5854;
    --forest:#2D4A3E; --forest-dim:#23392F; --gold:#C9A648; --gold-soft:#F1E7CC;
    --sage:#DDE5DE; --sage-line:#C7D2C8; --clay:#B8543E; --clay-soft:#F3E4E0;
    --rail:#E7E3DA; --shadow:0 1px 2px rgba(26,25,24,0.04), 0 4px 16px rgba(26,25,24,0.06);
    """

st.markdown(f"""
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,500;9..144,600;9..144,700&family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap" rel="stylesheet">
<style>
:root{{ {theme_vars} }}

/* ---------- BASE ---------- */
.stApp{{ background:var(--paper) !important; }}
html, body, [class*="css"]{{ font-family:'Inter',sans-serif; color:var(--ink); }}
h1,h2,h3{{ font-family:'Fraunces',serif !important; font-weight:600 !important; color:var(--ink) !important; letter-spacing:-0.01em; }}
p, span, label, div {{ color:var(--ink); }}
::selection{{ background:var(--gold-soft); }}
a{{ color:var(--forest); }}

/* Kill Streamlit chrome */
#MainMenu, footer, header[data-testid="stHeader"]{{ background:transparent; }}
.block-container{{ padding-top:2.2rem; max-width:1180px; }}

/* Inputs & Sidebar */
section[data-testid="stSidebar"]{{ background:var(--paper-raised) !important; border-right:1px solid var(--rail); }}
input, textarea, [data-baseweb="select"] > div {{
    background:var(--paper) !important; border:1px solid var(--rail) !important;
    border-radius:6px !important; color:var(--ink) !important;
    font-family:'JetBrains Mono',monospace !important; font-size:0.85rem !important;
}}
label p {{ font-size:0.75rem !important; font-weight:600 !important; color:var(--ink-soft) !important; text-transform:uppercase; letter-spacing:0.05em; }}

/* Tabs */
.stTabs [data-baseweb="tab-list"]{{ gap:26px; border-bottom:1px solid var(--rail); background:transparent; }}
.stTabs [data-baseweb="tab"]{{ background:transparent !important; color:var(--ink-soft) !important; font-weight:600; font-size:0.92rem; border-bottom:2px solid transparent !important; }}
.stTabs [aria-selected="true"]{{ color:var(--forest) !important; border-bottom:2px solid var(--forest) !important; }}

/* Buttons */
.stButton > button{{
    background:var(--forest) !important; color:#fff !important; border:none !important;
    border-radius:7px !important; font-weight:600 !important; font-size:0.88rem !important;
    padding:0.6rem 1.1rem !important; transition:background .15s ease, transform .08s ease !important;
}}
.stButton > button:hover{{ background:var(--forest-dim) !important; }}

/* Custom UI Components */
.eng-panel {{ background:var(--paper-raised); border:1px solid var(--rail); border-radius:10px; padding:20px; box-shadow:var(--shadow); margin-bottom:16px; }}
.eng-job-card {{ background:var(--paper); border:1px solid var(--rail); border-radius:8px; padding:16px; margin-bottom:8px; border-left: 4px solid var(--forest); }}
.eng-badge {{ font-family:'JetBrains Mono',monospace; font-size:0.65rem; background:var(--sage); color:var(--forest); padding:3px 8px; border-radius:5px; text-transform:uppercase; }}

/* Polished Resume Card */
.resume-view {{ background:var(--paper-raised); padding:30px; border-radius:10px; border:1px solid var(--rail); box-shadow:var(--shadow); }}
.resume-name {{ font-family:'Fraunces', serif; font-size:2rem; font-weight:700; color:var(--ink); margin-bottom:5px; }}
.resume-summary {{ font-size:0.9rem; color:var(--ink-soft); line-height:1.6; margin-bottom:20px; }}
.resume-section {{ font-family:'JetBrains Mono', monospace; font-size:0.8rem; color:var(--forest); text-transform:uppercase; border-bottom:1px solid var(--rail); padding-bottom:4px; margin-top:20px; margin-bottom:10px; letter-spacing:1px; }}
.resume-chip {{ display:inline-block; background:var(--sage); color:var(--forest); padding:4px 10px; border-radius:20px; font-size:0.75rem; font-weight:600; margin:3px; }}
.resume-body {{ font-size:0.85rem; color:var(--ink); line-height:1.6; white-space:pre-wrap; }}
</style>
""", unsafe_allow_html=True)

# ========== DATA & SIDEBAR FILTERS ================
GLOBAL_LOCATIONS = {
    "Asia": {"India": ["Delhi", "Mumbai", "Bangalore", "Pune", "Gurugram"], "Singapore": ["Singapore City"]},
    "North America": {"United States": ["New York", "San Francisco", "Remote"], "Canada": ["Toronto", "Vancouver"]},
    "Europe": {"United Kingdom": ["London"], "Germany": ["Berlin"]}
}
JOB_TAXONOMY = {
    "Tech & Engineering": ["Software Developer", "Python Developer", "Full-Stack Developer", "AI/ML Engineer"],
    "Business & Finance": ["Financial Analyst", "Business Operations Manager"],
    "Creative & Design": ["UI/UX Designer", "Graphic Designer"]
}
EMP_TYPES = ["Full-time", "Part-time", "Internship", "Freelance"]

with st.sidebar:
    with st.expander("🔑 API Keys", expanded=False):
        TAVILY_API_KEY = st.text_input("Tavily API key", type="password")
        GOOGLE_API_KEY = st.text_input("Gemini API key", type="password")

    st.markdown("### 🌍 Location")
    region = st.selectbox("Region", list(GLOBAL_LOCATIONS.keys()), index=0)
    country = st.selectbox("Country", list(GLOBAL_LOCATIONS[region].keys()), index=0)
    city = st.selectbox("City / State", GLOBAL_LOCATIONS[region][country], index=0)
    selected_location = f"{city}, {country}"

    st.markdown("### 💼 Job Preferences")
    selected_industries = st.multiselect("Industry", list(JOB_TAXONOMY.keys()), default=["Tech & Engineering"])
    available_roles = [role for ind in selected_industries for role in JOB_TAXONOMY[ind]]
    selected_roles = st.multiselect("Role", available_roles, default=["Software Developer"] if available_roles else [])
    emp_type = st.multiselect("Contract Type", EMP_TYPES, default=["Full-time"])

    st.markdown("### 📊 Advanced Filters")
    exp_years = st.slider("Experience (Years)", 0, 20, (0, 3))
    salary_range = st.slider("Salary Range (USD)", 0, 300000, (50000, 120000), step=5000)

def check_apis():
    return bool(TAVILY_API_KEY and GOOGLE_API_KEY)

def get_llm():
    return ChatGoogleGenerativeAI(model='gemini-1.5-flash', google_api_key=GOOGLE_API_KEY, temperature=0.3)

# ========== EXPORT GENERATORS ===================
def generate_docx(resume_json):
    doc = Document()
    doc.add_heading(resume_json.get("Name", "Resume"), 0)
    doc.add_paragraph(resume_json.get("Summary", ""))
    
    doc.add_heading("Skills", level=1)
    skills = resume_json.get("Skills", [])
    if isinstance(skills, list):
        doc.add_paragraph(", ".join(skills))
    else:
        doc.add_paragraph(str(skills))
        
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(resume_json.get("Experience", ""))
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph(resume_json.get("Education", ""))
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_pdf(resume_json):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt=resume_json.get("Name", "Resume"), ln=True, align='L')
    
    pdf.set_font("Arial", '', 11)
    pdf.multi_cell(0, 7, txt=resume_json.get("Summary", ""))
    
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, txt="Skills", ln=True)
    pdf.set_font("Arial", '', 11)
    skills = resume_json.get("Skills", [])
    pdf.multi_cell(0, 7, txt=", ".join(skills) if isinstance(skills, list) else str(skills))
    
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, txt="Experience", ln=True)
    pdf.set_font("Arial", '', 11)
    pdf.multi_cell(0, 7, txt=resume_json.get("Experience", ""))
    
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, txt="Education", ln=True)
    pdf.set_font("Arial", '', 11)
    pdf.multi_cell(0, 7, txt=resume_json.get("Education", ""))
    
    return pdf.output(dest='S').encode('latin1')

# ========== AI LOGIC ============================
def extract_text(resp):
    c = resp.content
    return "".join([b.get("text", "") if isinstance(b, dict) else str(b) for b in c]) if isinstance(c, list) else str(c)

def generate_base_resume(llm, raw_text, role):
    p = f"Convert this raw user data into structured JSON resume for a {role}. Format MUST be exactly: {{'Name': '..', 'Summary': '..', 'Skills': ['..'], 'Experience': '..', 'Education': '..'}}. Data: {raw_text}"
    t = extract_text(llm.invoke(p)).replace("```json", "").replace("```", "").strip()
    try: return json.loads(t)
    except: return {"Name": "Professional", "Summary": t, "Skills": [], "Experience": "See Summary", "Education": ""}

def tailor_resume(llm, base_json, jd_text):
    p = f"Tailor this base resume JSON to match this JD. Output ONLY valid JSON.\nBase: {json.dumps(base_json)}\nJD: {jd_text}"
    t = extract_text(llm.invoke(p)).replace("```json", "").replace("```", "").strip()
    try: return json.loads(t)
    except: return base_json

def score_ats(llm, resume, jd):
    return extract_text(llm.invoke(f"Act as strict ATS. Score 0-100. Find missing keywords and formatting issues.\nResume:{resume}\nJD:{jd}"))

def get_cold_email(llm, base, role, company):
    return extract_text(llm.invoke(f"Write a professional cold email applying for {role} at {company} using highlights from this resume: {json.dumps(base)}. Under 150 words."))

# ========== MAIN LAYOUT =========================
tab1, tab2, tab3, tab4 = st.tabs(["01 Resume Builder", "02 Job Board", "03 ATS Scorer", "04 Cold Emails"])

# --- TAB 1: RESUME BUILDER ---
with tab1:
    st.markdown("<h2>Build & Polish Your Profile</h2>", unsafe_allow_html=True)
    colA, colB = st.columns([1, 1.2])
    
    with colA:
        st.markdown("<div class='eng-panel'>", unsafe_allow_html=True)
        roles_str = ", ".join(selected_roles) if selected_roles else "Professional"
        target_role = st.text_input("Target Roles", value=roles_str)
        raw_exp = st.text_area("Raw Experience & Details", height=200, value="BCA student, skilled in Python, PHP, web design. Building robust tech solutions.")
        
        if st.button("Generate Master Resume", use_container_width=True):
            if check_apis():
                with st.spinner("Structuring your AI resume..."):
                    st.session_state.base_resume = generate_base_resume(get_llm(), raw_exp, target_role)
            else: st.error("Add API Keys in Sidebar")
        st.markdown("</div>", unsafe_allow_html=True)
        
    with colB:
        if st.session_state.base_resume:
            r = st.session_state.base_resume
            # Beautiful HTML Render (No JSON)
            skills = r.get("Skills", [])
            chips = "".join([f"<span class='resume-chip'>{s}</span>" for s in skills]) if isinstance(skills, list) else f"<span class='resume-chip'>{skills}</span>"
            
            st.markdown(f"""
            <div class='resume-view'>
                <div class='resume-name'>{r.get('Name', 'Your Name')}</div>
                <div class='resume-summary'>{r.get('Summary', '')}</div>
                <div class='resume-section'>Skills & Competencies</div>
                <div>{chips}</div>
                <div class='resume-section'>Professional Experience</div>
                <div class='resume-body'>{r.get('Experience', '')}</div>
                <div class='resume-section'>Education</div>
                <div class='resume-body'>{r.get('Education', '')}</div>
            </div>
            """, unsafe_allow_html=True)
            
            # Export Buttons
            st.markdown("<br>", unsafe_allow_html=True)
            d_col1, d_col2 = st.columns(2)
            d_col1.download_button("📥 Download PDF", data=generate_pdf(r), file_name="Resume.pdf", mime="application/pdf", use_container_width=True)
            d_col2.download_button("📥 Download Word (DOCX)", data=generate_docx(r), file_name="Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        else:
            st.info("Your polished, ATS-friendly resume preview will appear here.")

# --- TAB 2: JOB BOARD ---
with tab2:
    st.markdown("<h2>Live Web Job Scraper</h2>", unsafe_allow_html=True)
    display_roles = ", ".join(selected_roles) if selected_roles else "Any role"
    
    if st.button("↻ Scrape Web Now (Tavily)", use_container_width=True):
        if not selected_roles or not emp_type:
            st.warning("Select roles and employment types in the sidebar.")
        elif check_apis():
            with st.spinner(f"Scraping roles in {selected_location}..."):
                client = TavilyClient(api_key=TAVILY_API_KEY)
                jobs = []
                for e_type in emp_type:
                    query = f"Latest {e_type} jobs for {display_roles} in {selected_location} {exp_years[0]}-{exp_years[1]} years experience salary ${salary_range[0]} apply"
                    res = client.search(query, search_depth="advanced", max_results=4)
                    for r in res.get("results", []):
                        jobs.append({"title": r.get("title", "Unknown"), "url": r.get("url", "#"), "desc": r.get("content", ""), "type": e_type, "company": r.get("title", "Company").split("-")[0].strip()})
                st.session_state.jobs_data = jobs
                st.success(f"Found {len(jobs)} live postings.")
        else: st.error("Add API keys first.")

    # Render Jobs with Individual & Bulk Apply
    if st.session_state.jobs_data:
        st.markdown("---")
        selected_for_apply = []
        
        for idx, job in enumerate(st.session_state.jobs_data):
            st.markdown(f"""
            <div class='eng-job-card'>
                <div style="display:flex; justify-content:space-between;">
                    <div><b style="font-size:1.1rem; color:var(--ink);">{job['title']}</b> <span style="color:var(--ink-soft); font-size:0.9rem;">at {job['company']}</span></div>
                    <span class='eng-badge'>{job['type']}</span>
                </div>
                <p style="font-size:0.85rem; color:var(--ink-soft); margin:8px 0;">{job['desc'][:250]}...</p>
                <a href="{job['url']}" target="_blank" style="font-size:0.8rem;">View Original Posting ↗</a>
            </div>
            """, unsafe_allow_html=True)
            
            # Action Buttons Row
            act1, act2, act3 = st.columns([1.5, 2, 5])
            if act1.button(f"⚡ Apply Now", key=f"app_single_{idx}"):
                if not st.session_state.base_resume: st.error("Build Base Resume First!")
                else:
                    with st.spinner("Tailoring..."):
                        st.session_state.tailored_resumes[job['title']] = tailor_resume(get_llm(), st.session_state.base_resume, job['desc'])
                        st.success(f"Applied to {job['title']}! (Simulated)")
            if act2.checkbox("Select for Bulk Apply", key=f"sel_{idx}"):
                selected_for_apply.append(job)
            st.markdown("<br>", unsafe_allow_html=True)

        if selected_for_apply:
            st.markdown(f"**{len(selected_for_apply)} roles selected for bulk submission.**")
            if st.button("🚀 Apply to All Selected", use_container_width=True):
                if not st.session_state.base_resume: st.error("Build Base Resume First!")
                else:
                    for j in selected_for_apply:
                        st.session_state.tailored_resumes[j['title']] = tailor_resume(get_llm(), st.session_state.base_resume, j['desc'])
                        st.success(f"Successfully processed {j['title']}")

# --- TAB 3: ATS SCORER ---
with tab3:
    st.markdown("<h2>Resume ATS Scorer</h2>", unsafe_allow_html=True)
    st.markdown("<div class='eng-panel'>", unsafe_allow_html=True)
    mode = st.radio("Resume Source", ["Use Current AI Resume", "Upload Custom PDF"], horizontal=True)
    target_jd = st.text_area("Target Job Description (Optional)", height=100)
    
    res_text = ""
    if mode == "Upload Custom PDF":
        pdf = st.file_uploader("Upload PDF", type=["pdf"])
        if pdf: res_text = PyPDF2.PdfReader(pdf).pages[0].extract_text()
    else:
        res_text = json.dumps(st.session_state.base_resume) if st.session_state.base_resume else ""
        
    if st.button("Score Resume", use_container_width=True):
        if not res_text: st.warning("Provide a resume first.")
        elif check_apis():
            with st.spinner("Scoring against ATS logic..."):
                st.markdown(score_ats(get_llm(), res_text, target_jd or "General Industry Standards"))
    st.markdown("</div>", unsafe_allow_html=True)

# --- TAB 4: COLD EMAILS ---
with tab4:
    st.markdown("<h2>AI Cold Email Generator</h2>", unsafe_allow_html=True)
    if st.session_state.jobs_data and st.session_state.base_resume:
        st.markdown("<div class='eng-panel'>", unsafe_allow_html=True)
        titles = [j['title'] for j in st.session_state.jobs_data]
        sel = st.selectbox("Select target role", titles)
        j = next(j for j in st.session_state.jobs_data if j['title'] == sel)
        
        if st.button("Draft Cold Email", use_container_width=True):
            if check_apis():
                with st.spinner("Drafting..."):
                    st.session_state.draft = get_cold_email(get_llm(), st.session_state.base_resume, j['title'], j['company'])
        
        if "draft" in st.session_state:
            email = st.text_area("Review Email", value=st.session_state.draft, height=200)
            if st.button("Send Email (Simulated)", use_container_width=True):
                st.success(f"Email sent to recruiter at {j['company']}!")
        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.info("Complete Step 1 and Step 2 to unlock tailored cold emails.")
